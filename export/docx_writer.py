from __future__ import annotations

from typing import List
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from context.placeholder_manager import split_placeholder_segments
from models.aha import AhaDoc
from models.csp import CspDoc


def _add_paragraph_with_placeholders(doc: Document, text: str, style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for segment, is_placeholder in split_placeholder_segments(text or ""):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = bold or is_placeholder
        if is_placeholder:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return paragraph


def _add_cover_page(doc: Document, csp: CspDoc) -> None:
    """Add clean title block - no duplicates, no wrap-around."""
    from datetime import datetime
    doc.add_heading('Comprehensive Site-Specific Construction Safety and Health Plan', level=0)
    doc.add_paragraph()  # Spacing

    # Title block fields (fail if missing per validation)
    _add_paragraph_with_placeholders(doc, f"Project: {csp.project_name or '[Insert Project Name]'}")
    if csp.project_number:
        _add_paragraph_with_placeholders(doc, f"Project Number: {csp.project_number}")
    _add_paragraph_with_placeholders(doc, f"Location: {csp.location or '[Insert Location]'}")
    if csp.owner:
        _add_paragraph_with_placeholders(doc, f"Owner: {csp.owner}")
    if csp.general_contractor:
        _add_paragraph_with_placeholders(doc, f"Prime Contractor: {csp.general_contractor}")

    doc.add_paragraph()  # Spacing

    doc.add_paragraph(f"Document ID: CSP-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
    doc.add_paragraph(f"Version: 1.0")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_page_break()


def _add_revision_log(doc: Document) -> None:
    doc.add_heading('Revision Log', level=1)
    table = doc.add_table(rows=2, cols=4)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    hdr = table.rows[0].cells
    hdr[0].text = 'Revision'
    hdr[1].text = 'Date'
    hdr[2].text = 'Description'
    hdr[3].text = 'Prepared By'
    row = table.rows[1].cells
    row[0].text = '0'
    row[1].text = ''
    row[2].text = 'Initial issue'
    row[3].text = ''
    doc.add_paragraph()


def write_aha_book(docs: List[AhaDoc], output_path: str) -> str:
    """Write a simple AHA Book DOCX with headings, tables, and citations."""
    doc = Document()

    # Title
    doc.add_heading('Activity Hazard Analysis (AHA) Book', level=0)

    # Index
    doc.add_heading('Index', level=1)
    for aha in docs:
        p = doc.add_paragraph()
        run = p.add_run(aha.name)
        run.bold = True

    for aha in docs:
        doc.add_page_break()
        doc.add_heading(aha.name, level=1)
        doc.add_paragraph(f"Activity: {aha.activity}")
        if aha.hazards:
            doc.add_paragraph("Hazards:")
            for h in aha.hazards:
                doc.add_paragraph(h, style='List Bullet')

        # Steps table
        doc.add_heading('Work Sequence, Hazards, Controls, PPE, Permits/Training', level=2)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = 'Step'
        hdr[1].text = 'Hazards'
        hdr[2].text = 'Controls'
        hdr[3].text = 'PPE'
        hdr[4].text = 'Permits/Training'
        for it in aha.items:
            row = table.add_row().cells
            row[0].text = it.step
            row[1].text = "\n".join(it.hazards)
            row[2].text = "\n".join(it.controls)
            row[3].text = "\n".join(it.ppe)
            row[4].text = "\n".join(it.permits_training)

        # Citations
        if aha.citations:
            doc.add_heading('Citations', level=2)
            for c in aha.citations:
                line = f"§ {c.section_path} | page {c.page_label or c.page_number or ''} — \"{c.quote_anchor}\""
                doc.add_paragraph(line, style='List Number')

        # Codes Covered
        if getattr(aha, 'codes_covered', []):
            doc.add_heading('Codes Covered', level=2)
            for code in aha.codes_covered:
                doc.add_paragraph(code, style='List Bullet')

        # PPE table
        ppe_items: list[str] = []
        permits_items: list[str] = []
        for it in aha.items:
            ppe_items.extend(it.ppe)
            permits_items.extend(it.permits_training)
        if ppe_items:
            doc.add_heading('PPE', level=2)
            t = doc.add_table(rows=1, cols=2)
            hdr = t.rows[0].cells
            hdr[0].text = 'Item'
            hdr[1].text = 'Specification'
            for p in ppe_items:
                row = t.add_row().cells
                row[0].text = p
                row[1].text = ''
        if permits_items:
            doc.add_heading('Permits & Training', level=2)
            for p in permits_items:
                doc.add_paragraph(p, style='List Bullet')

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def write_aha_single(aha: AhaDoc, output_path: str) -> str:
    doc = Document()
    doc.add_heading(aha.name, level=0)
    doc.add_paragraph(f"Activity: {aha.activity}")
    if aha.hazards:
        doc.add_paragraph("Hazards:")
        for h in aha.hazards:
            doc.add_paragraph(h, style='List Bullet')

    doc.add_heading('Work Sequence, Hazards, Controls, PPE, Permits/Training', level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Hazards'
    hdr[2].text = 'Controls'
    hdr[3].text = 'PPE'
    hdr[4].text = 'Permits/Training'
    for it in aha.items:
        row = table.add_row().cells
        row[0].text = it.step
        row[1].text = "\n".join(it.hazards)
        row[2].text = "\n".join(it.controls)
        row[3].text = "\n".join(it.ppe)
        row[4].text = "\n".join(it.permits_training)

    if aha.citations:
        doc.add_heading('Citations', level=2)
        for c in aha.citations:
            line = f"§ {c.section_path} | page {c.page_label or c.page_number or ''} — \"{c.quote_anchor}\""
            doc.add_paragraph(line, style='List Number')

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def write_csp_docx(csp: CspDoc, output_path: str) -> str:
    doc = Document()
    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass

    _add_cover_page(doc, csp)
    _add_revision_log(doc)

    for idx, sec in enumerate(csp.sections):
        if idx == 0:
            doc.add_page_break()
        doc.add_heading(sec.name, level=1)
        for p in sec.paragraphs:
            text = (p or '').strip()
            if not text:
                continue
            if text.startswith('- '):
                _add_paragraph_with_placeholders(doc, text[2:], style='List Bullet')
            elif text.startswith('• '):
                _add_paragraph_with_placeholders(doc, text[2:], style='List Bullet')
            elif text.startswith('•'):
                _add_paragraph_with_placeholders(doc, text[1:].strip(), style='List Bullet')
            elif text.startswith('References:'):
                _add_paragraph_with_placeholders(doc, text, bold=True)
            else:
                _add_paragraph_with_placeholders(doc, text)
        if sec.citations:
            doc.add_heading('Citations', level=2)
            for c in sec.citations:
                pg = c.page_label or c.page_number or ''
                citation_text = f"{c.section_path}"
                if pg:
                    citation_text += f" | page {pg}"
                if c.quote_anchor:
                    citation_text += f" — \"{c.quote_anchor}\""
                _add_paragraph_with_placeholders(doc, citation_text, style='List Number')

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


