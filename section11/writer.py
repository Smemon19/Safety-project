"""Artifact writers for the Section 11 Generator."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, List

from docx import Document  # type: ignore

from section11.models import (
    CategoryBundle,
    CategoryStatus,
    ComplianceMatrixRow,
    Section11Artifacts,
    Section11Run,
)


def _status_badge(status: CategoryStatus) -> str:
    return status.value


def _write_compliance_matrix(rows: List[ComplianceMatrixRow]) -> List[str]:
    lines = ["# Section 11.0 Compliance Matrix", ""]
    lines.append("| Category | Codes | AHA Status | Safety Plan Status | Project Evidence | EM Evidence |")
    lines.append("| --- | --- | --- | --- | --- | --- |")
    for row in rows:
        codes = "<br>".join(row.codes) if row.codes else "—"
        lines.append(
            "| {category} | {codes} | {aha} | {plan} | {proj} | {em} |".format(
                category=row.category,
                codes=codes,
                aha=_status_badge(row.aha_status),
                plan=_status_badge(row.plan_status),
                proj=row.project_evidence_count,
                em=row.em_evidence_count,
            )
        )
    lines.append("")
    return lines


def _write_bundle_markdown(bundle: CategoryBundle) -> List[str]:
    lines = [f"## {bundle.category}", ""]
    lines.append("### Activity Hazard Analysis (Hazards Only)")
    lines.append("")
    if bundle.aha.hazards:
        lines.append("**Hazards Identified:**")
        for hazard in bundle.aha.hazards:
            lines.append(f"- {hazard}")
        lines.append("")
    if bundle.aha.narrative:
        lines.append("**Hazard Narrative:**")
        for paragraph in bundle.aha.narrative:
            lines.append(f"- {paragraph}")
        lines.append("")
    if bundle.aha.citations:
        lines.append("**AHA Citations:**")
        for citation in bundle.aha.citations:
            section = citation.get("section_path", "")
            page = citation.get("page_label") or citation.get("page_number") or ""
            if page:
                lines.append(f"- § {section} (p. {page})")
            else:
                lines.append(f"- § {section}")
        lines.append("")
    if bundle.aha.pending_reason:
        lines.append(f"_Status: {bundle.aha.pending_reason}_")
        lines.append("")

    lines.append("### Safety Plan (Controls Only)")
    lines.append("")
    if bundle.plan.controls:
        lines.append("**Controls and Procedures:**")
        for control in bundle.plan.controls:
            lines.append(f"- {control}")
        lines.append("")
    if bundle.plan.ppe:
        lines.append("**PPE:**")
        for item in bundle.plan.ppe:
            lines.append(f"- {item}")
        lines.append("")
    if bundle.plan.permits:
        lines.append("**Permits / Training / Inspections:**")
        for permit in bundle.plan.permits:
            lines.append(f"- {permit}")
        lines.append("")
    if bundle.plan.citations:
        lines.append("**Safety Plan Citations:**")
        for citation in bundle.plan.citations:
            section = citation.get("section_path", "")
            page = citation.get("page_label") or citation.get("page_number") or ""
            if page:
                lines.append(f"- § {section} (p. {page})")
            else:
                lines.append(f"- § {section}")
        lines.append("")
    if bundle.plan.pending_reason:
        lines.append(f"_Status: {bundle.plan.pending_reason}_")
        lines.append("")
    return lines


def write_section11_markdown(base_dir: Path, bundles: List[CategoryBundle], matrix: List[ComplianceMatrixRow]) -> Path:
    base_dir.mkdir(parents=True, exist_ok=True)
    path = base_dir / "section11.md"
    lines = _write_compliance_matrix(matrix)
    lines.append("---")
    for bundle in bundles:
        lines.extend(_write_bundle_markdown(bundle))
        lines.append("---")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_bundle_markdown(base_dir: Path, bundle: CategoryBundle, prefix: str) -> Path:
    safe_category = bundle.category.lower().replace(" ", "_").replace("/", "-")
    path = base_dir / f"{prefix}_{safe_category}.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(_write_bundle_markdown(bundle)), encoding="utf-8")
    return path


def write_section11_docx(base_dir: Path, bundles: List[CategoryBundle], matrix: List[ComplianceMatrixRow]) -> Path:
    document = Document()
    document.add_heading("Section 11 Safety Plan", level=1)
    document.add_heading("Section 11.0 Compliance Matrix", level=2)
    table = document.add_table(rows=1, cols=6)
    headers = ["Category", "Codes", "AHA Status", "Safety Plan Status", "Project Evidence", "EM Evidence"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in matrix:
        cells = table.add_row().cells
        cells[0].text = row.category
        cells[1].text = "\n".join(row.codes) if row.codes else "—"
        cells[2].text = _status_badge(row.aha_status)
        cells[3].text = _status_badge(row.plan_status)
        cells[4].text = str(row.project_evidence_count)
        cells[5].text = str(row.em_evidence_count)

    for bundle in bundles:
        document.add_page_break()
        document.add_heading(bundle.category, level=2)
        document.add_heading("Activity Hazard Analysis", level=3)
        if bundle.aha.hazards:
            document.add_paragraph("Hazards Identified:")
            for hazard in bundle.aha.hazards:
                document.add_paragraph(hazard, style="List Bullet")
        if bundle.aha.narrative:
            document.add_paragraph("Hazard Narrative:")
            for paragraph in bundle.aha.narrative:
                document.add_paragraph(paragraph, style="List Bullet")
        if bundle.aha.citations:
            document.add_paragraph("AHA Citations:")
            for citation in bundle.aha.citations:
                section = citation.get("section_path", "")
                page = citation.get("page_label") or citation.get("page_number") or ""
                document.add_paragraph(f"§ {section} (p. {page})", style="List Bullet")
        if bundle.aha.pending_reason:
            document.add_paragraph(bundle.aha.pending_reason)

        document.add_heading("Safety Plan", level=3)
        if bundle.plan.controls:
            document.add_paragraph("Controls and Procedures:")
            for control in bundle.plan.controls:
                document.add_paragraph(control, style="List Bullet")
        if bundle.plan.ppe:
            document.add_paragraph("PPE:")
            for item in bundle.plan.ppe:
                document.add_paragraph(item, style="List Bullet")
        if bundle.plan.permits:
            document.add_paragraph("Permits / Training / Inspections:")
            for permit in bundle.plan.permits:
                document.add_paragraph(permit, style="List Bullet")
        if bundle.plan.citations:
            document.add_paragraph("Safety Plan Citations:")
            for citation in bundle.plan.citations:
                section = citation.get("section_path", "")
                page = citation.get("page_label") or citation.get("page_number") or ""
                document.add_paragraph(f"§ {section} (p. {page})", style="List Bullet")
        if bundle.plan.pending_reason:
            document.add_paragraph(bundle.plan.pending_reason)

    base_dir.mkdir(parents=True, exist_ok=True)
    path = base_dir / "section11.docx"
    document.save(str(path))
    return path


def write_section11_json(base_dir: Path, run: Section11Run) -> Path:
    base_dir.mkdir(parents=True, exist_ok=True)
    
    # Serialize with enum values instead of enum objects
    def serialize_status(status: CategoryStatus) -> str:
        return status.value
    
    def serialize_matrix_row(row: ComplianceMatrixRow) -> Dict[str, object]:
        return {
            "category": row.category,
            "codes": row.codes,
            "aha_status": serialize_status(row.aha_status),
            "plan_status": serialize_status(row.plan_status),
            "project_evidence_count": row.project_evidence_count,
            "em_evidence_count": row.em_evidence_count,
            "aha_link": row.aha_link,
            "plan_link": row.plan_link,
        }
    
    def serialize_aha(aha) -> Dict[str, object]:
        return {
            "hazards": aha.hazards,
            "narrative": aha.narrative,
            "citations": aha.citations,
            "status": serialize_status(aha.status),
            "pending_reason": aha.pending_reason,
        }
    
    def serialize_plan(plan) -> Dict[str, object]:
        return {
            "controls": plan.controls,
            "ppe": plan.ppe,
            "permits": plan.permits,
            "citations": plan.citations,
            "project_evidence": plan.project_evidence,
            "em_evidence": plan.em_evidence,
            "status": serialize_status(plan.status),
            "pending_reason": plan.pending_reason,
        }
    
    payload: Dict[str, object] = {
        "run_id": run.run_id,
        "source_file": run.source_file.name,
        "matrix": [serialize_matrix_row(row) for row in run.matrix],
        "categories": [
            {
                "category": bundle.category,
                "codes": bundle.codes,
                "aha": serialize_aha(bundle.aha),
                "plan": serialize_plan(bundle.plan),
            }
            for bundle in run.bundles
        ],
    }
    path = base_dir / "section11_report.json"
    path.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
    return path


def allocate_artifacts(base_dir: Path) -> Section11Artifacts:
    manifest_dir = base_dir / "section11_bundle"
    manifest_dir.mkdir(parents=True, exist_ok=True)
    return Section11Artifacts(
        base_dir=base_dir,
        manifest_path=manifest_dir / "manifest.json",
        markdown_path=base_dir / "section11.md",
        docx_path=base_dir / "section11.docx",
        json_report_path=base_dir / "section11_report.json",
    )

